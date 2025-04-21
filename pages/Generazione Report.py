from docx import Document
from docx.shared import Inches
import plotly.graph_objects as go
import pandas as pd
import numpy as np
import io

# Dati simulati come nella dashboard
criteri = ["Accessibilità del verde", "Biodiversità", "Manutenzione e pulizia", "Funzione sociale", "Funzione ambientale"]
riepilogo = pd.DataFrame({
    "Parco": ["Parco Suardi", "Parco della Trucca", "Parco Goisis", "Parco Locatelli"],
    "Accessibilità del verde": [3, 5, 4, 3],
    "Biodiversità": [4, 5, 3, 2],
    "Manutenzione e pulizia": [3, 4, 2, 4],
    "Funzione sociale": [4, 4, 3, 2],
    "Funzione ambientale": [4, 5, 3, 3],
    "Punteggio Finale": [3.8, 4.8, 3.1, 2.8]
})

# Crea documento Word
doc = Document()
doc.add_heading("Report finale - Valutazione dei parchi urbani", 0)

# Riepilogo tabellare
doc.add_heading("📋 Riepilogo tabellare dei punteggi", level=1)
table = doc.add_table(rows=1, cols=len(riepilogo.columns))
hdr_cells = table.rows[0].cells
for i, col in enumerate(riepilogo.columns):
    hdr_cells[i].text = col

for _, row in riepilogo.iterrows():
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = str(round(val, 2)) if isinstance(val, float) else str(val)

# Radar chart
fig = go.Figure()
for _, row in riepilogo.iterrows():
    fig.add_trace(go.Scatterpolar(
        r=[row[c] for c in criteri],
        theta=criteri,
        fill='none',
        name=row["Parco"]
    ))
fig.update_layout(
    polar=dict(radialaxis=dict(visible=True, range=[0, 5])),
    showlegend=True
)

# Esporta immagine radar
img_bytes = fig.to_image(format="png", width=800, height=600, engine="kaleido")
img_stream = io.BytesIO(img_bytes)

doc.add_page_break()
doc.add_heading("📈 Grafico radar comparativo", level=1)
doc.add_paragraph("Il grafico seguente rappresenta la valutazione media per ciascun criterio nei diversi parchi.")
doc.add_picture(img_stream, width=Inches(6))

# Mappa descrizione
doc.add_page_break()
doc.add_heading("🗺 Descrizione della mappa dei punteggi", level=1)
doc.add_paragraph("La mappa interattiva mostra i punteggi geolocalizzati dei parchi. Colori: rosso (≤5), giallo (5-9), verde (>9).")

# Salva documento
output = io.BytesIO()
doc.save(output)
output.seek(0)
output.name = "Report_Finale_Valutazione_Parchi.docx"
output
