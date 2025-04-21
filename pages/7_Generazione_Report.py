import pandas as pd
from docx import Document
from docx.shared import Inches
import plotly.graph_objects as go
import io

# Simula il DataFrame `riepilogo` da Streamlit
criteri = ["Accessibilità del verde", "Biodiversità", "Manutenzione e pulizia", "Funzione sociale", "Funzione ambientale"]
riepilogo = pd.DataFrame({
    "Parco": ["Parco Suardi", "Parco della Trucca", "Parco Goisis", "Parco Locatelli"],
    "Accessibilità del verde": [3, 5, 4, 3],
    "Biodiversità": [4, 5, 3, 2],
    "Manutenzione e pulizia": [3, 4, 2, 4],
    "Funzione sociale": [4, 4, 3, 2],
    "Funzione ambientale": [4, 5, 3, 3],
    "Punteggio Finale": [3.8, 4.8, 3.1, 2.8]
})

# Crea documento Word
doc = Document()
doc.add_heading("Report finale - Valutazione dei parchi urbani", 0)

# Aggiungi tabella riepilogo
doc.add_heading("📋 Riepilogo tabellare", level=1)
table = doc.add_table(rows=1, cols=len(riepilogo.columns))
for i, col in enumerate(riepilogo.columns):
    table.rows[0].cells[i].text = col
for _, row in riepilogo.iterrows():
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = str(round(val, 2)) if isinstance(val, float) else str(val)

# Aggiungi grafico radar
fig = go.Figure()
for _, row in riepilogo.iterrows():
    fig.add_trace(go.Scatterpolar(
        r=[row[c] for c in criteri],
        theta=criteri,
        fill='none',
        name=row["Parco"]
    ))
fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 5])), showlegend=True)
img_bytes = fig.to_image(format='png', engine='kaleido')
doc.add_page_break()
doc.add_heading("📈 Grafico radar comparativo", level=1)
doc.add_paragraph("Il grafico rappresenta le performance dei parchi per ciascun criterio.")
doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))

# Conclusione
doc.add_page_break()
doc.add_heading("🗺 Descrizione della mappa", level=1)
doc.add_paragraph("La mappa interattiva mostra i punteggi geolocalizzati. Rosso ≤5, Giallo 5–9, Verde >9.")

# Salva
doc.save("Report_Finale_Valutazione_Parchi.docx")
